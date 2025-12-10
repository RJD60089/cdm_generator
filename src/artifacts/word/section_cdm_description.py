# src/artifacts/word/section_cdm_description.py
"""CDM Description section for Word CDM document."""

from docx import Document
from src.artifacts.common.cdm_extractor import CDMExtractor


def add_cdm_description(doc: Document, extractor: CDMExtractor):
    """Add CDM Description section."""
    
    doc.add_heading("2. CDM Description", level=1)
    
    # Domain description
    doc.add_heading("Domain Overview", level=2)
    
    description = extractor.domain_description or f"The {extractor.domain} domain."
    doc.add_paragraph(description)
    
    # Purpose
    doc.add_heading("Purpose", level=2)
    
    doc.add_paragraph(
        f"The {extractor.domain} CDM serves as the authoritative definition for "
        f"{extractor.domain.lower()}-related data elements within the organization. It enables:"
    )
    
    purposes = [
        "Consistent data integration across source systems",
        "Standardized reporting and analytics",
        "Clear data governance and ownership",
        "Interoperability with external partners and standards",
        "Foundation for data quality management"
    ]
    
    for purpose in purposes:
        p = doc.add_paragraph(purpose, style='List Bullet')
    
    # Scope
    doc.add_heading("Scope", level=2)
    
    entities = extractor.get_entities()
    entity_names = [e.name for e in entities]
    
    doc.add_paragraph(
        f"This CDM covers {len(entities)} entities: {', '.join(entity_names)}. "
        f"It defines the canonical structure, data types, relationships, and business rules "
        f"for each entity."
    )
    
    doc.add_page_break()