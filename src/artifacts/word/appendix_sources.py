# src/artifacts/word/appendix_sources.py
"""Source Files Appendix for Word CDM document."""

from docx import Document
from src.config.config_parser import AppConfig
from src.artifacts.common.cdm_extractor import CDMExtractor


def add_sources_appendix(doc: Document, extractor: CDMExtractor, config: AppConfig):
    """Add Source Files appendix."""
    
    doc.add_heading("Appendix C: Source Files", level=1)
    
    doc.add_paragraph(
        "The following source files were used to generate this CDM."
    )
    
    # Configuration info
    doc.add_heading("Configuration", level=2)
    
    config_info = doc.add_paragraph()
    config_info.add_run("Domain: ").bold = True
    config_info.add_run(f"{config.cdm.domain}\n")
    config_info.add_run("Type: ").bold = True
    config_info.add_run(f"{config.cdm.type}\n")
    config_info.add_run("Output Directory: ").bold = True
    config_info.add_run(f"{config.output.directory}\n")
    
    # Source files by type
    doc.add_heading("Input Files by Source Type", level=2)
    
    has_sources = False
    
    # Guardrails files (List[str])
    if config.has_guardrails():
        has_sources = True
        doc.add_heading("Guardrails", level=3)
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells[0]
        hdr.text = "File Path"
        for paragraph in hdr.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        
        for gfile in config.guardrails:
            row = table.add_row().cells
            row[0].text = str(gfile)
        
        doc.add_paragraph()
    
    # Glue files (List[str])
    if config.has_glue():
        has_sources = True
        doc.add_heading("Glue (AWS)", level=3)
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells[0]
        hdr.text = "File Path"
        for paragraph in hdr.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        
        for glue_file in config.glue:
            row = table.add_row().cells
            row[0].text = str(glue_file)
        
        doc.add_paragraph()
    
    # NCPDP standards (List[Dict] with name, version, file, etc.)
    if config.has_ncpdp():
        has_sources = True
        doc.add_heading("NCPDP Standards", level=3)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Standard"
        hdr[1].text = "Version"
        hdr[2].text = "File"
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # General standards
        for standard in config.ncpdp_general_standards:
            row = table.add_row().cells
            row[0].text = standard.get('name', '-')
            row[1].text = standard.get('version', '-')
            row[2].text = standard.get('file', '-')
        
        # Script standards
        for standard in config.ncpdp_script_standards:
            row = table.add_row().cells
            row[0].text = standard.get('name', '-')
            row[1].text = standard.get('version', '-')
            row[2].text = standard.get('file', '-')
        
        doc.add_paragraph()
    
    # FHIR resources (List[Dict] with file, file_type, resource_type, etc.)
    if config.has_fhir():
        has_sources = True
        doc.add_heading("FHIR Resources", level=3)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Resource Type"
        hdr[1].text = "File Type"
        hdr[2].text = "File"
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for resource in config.fhir_igs:
            row = table.add_row().cells
            row[0].text = resource.get('resource_type', resource.get('name', '-'))
            row[1].text = resource.get('file_type', '-')
            row[2].text = resource.get('file', '-')
        
        doc.add_paragraph()
    
    # DDL files (List[str])
    if config.ddl:
        has_sources = True
        doc.add_heading("DDL Files", level=3)
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells[0]
        hdr.text = "File Path"
        for paragraph in hdr.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        
        for ddl_file in config.ddl:
            row = table.add_row().cells
            row[0].text = str(ddl_file)
        
        doc.add_paragraph()
    
    # Naming standard files (List[str])
    if config.naming_standard:
        has_sources = True
        doc.add_heading("Naming Standards", level=3)
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells[0]
        hdr.text = "File Path"
        for paragraph in hdr.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        
        for ns_file in config.naming_standard:
            row = table.add_row().cells
            row[0].text = str(ns_file)
        
        doc.add_paragraph()
    
    if not has_sources:
        doc.add_paragraph("No source files configured.")