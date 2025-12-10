# src/artifacts/word/section_methodology.py
"""Methodology section for Word CDM document."""

from docx import Document
from src.artifacts.common.cdm_extractor import CDMExtractor


def add_methodology(doc: Document, extractor: CDMExtractor):
    """Add Methodology section."""
    
    doc.add_heading("3. Methodology", level=1)
    
    # Approach
    doc.add_heading("Approach", level=2)
    
    doc.add_paragraph(
        "The CDM was developed by analyzing and reconciling data definitions from multiple "
        "authoritative sources. Each source provides unique perspective on the data model:"
    )
    
    sources_desc = [
        ("Guardrails", "Internal business requirements and data governance rules defining "
                      "expected data elements and business logic."),
        ("Glue", "Existing operational data structures from the current system implementation, "
                "representing the as-is state."),
        ("NCPDP", "Industry standard definitions from the National Council for Prescription "
                 "Drug Programs, ensuring regulatory compliance."),
        ("FHIR", "Healthcare interoperability standards from HL7 FHIR, enabling integration "
                "with external healthcare systems.")
    ]
    
    for source, desc in sources_desc:
        p = doc.add_paragraph()
        p.add_run(f"{source}: ").bold = True
        p.add_run(desc)
    
    # Source priority
    doc.add_heading("Source Priority", level=2)
    
    doc.add_paragraph(
        "When sources provide conflicting definitions, the following priority order is applied:"
    )
    
    doc.add_paragraph("Guardrails → Glue → NCPDP → FHIR", style='Intense Quote')
    
    doc.add_paragraph(
        "This prioritization reflects the organization's approach of balancing internal "
        "business requirements with industry standards. Internal sources (Guardrails, Glue) "
        "take precedence as they represent specific business needs, while external standards "
        "(NCPDP, FHIR) provide reference definitions."
    )
    
    # Validation approach
    doc.add_heading("Validation Approach", level=2)
    
    doc.add_paragraph(
        "This CDM document is provided for SME review and validation. The following items "
        "require attention:"
    )
    
    validation_items = [
        "Critical Data Elements (Section 5) - Confirm business criticality designations",
        "Requires Review (Section 6) - Resolve flagged data quality and mapping questions",
        "SME Questions (Section 7) - Provide answers to outstanding questions",
        "Unmapped Fields (Section 8) - Determine disposition of source fields not in CDM"
    ]
    
    for item in validation_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
