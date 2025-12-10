# src/artifacts/word/section_entities.py
"""Entity Summaries section for Word CDM document."""

from docx import Document
from docx.shared import Pt, Inches
from src.artifacts.common.cdm_extractor import CDMExtractor


def add_entities_section(doc: Document, extractor: CDMExtractor):
    """Add Entity Summaries section."""
    
    doc.add_heading("4. Entity Summaries", level=1)
    
    doc.add_paragraph(
        f"This section provides an overview of the {extractor.entity_count} entities "
        f"defined in the {extractor.domain} CDM."
    )
    
    # Summary table
    doc.add_heading("Entity Overview", level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header
    headers = ["Entity", "Description", "Primary Key", "Attributes"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    entities = extractor.get_entities()
    
    for entity in entities:
        row_cells = table.add_row().cells
        row_cells[0].text = entity.name
        row_cells[1].text = (entity.description[:80] + "...") if len(entity.description) > 80 else entity.description
        row_cells[2].text = ", ".join(entity.primary_keys) if entity.primary_keys else "-"
        row_cells[3].text = str(entity.attribute_count)
    
    doc.add_paragraph()  # Spacing
    
    # Detailed entity sections
    doc.add_heading("Entity Details", level=2)
    
    for entity in entities:
        doc.add_heading(entity.name, level=3)
        
        # Description
        doc.add_paragraph(entity.description)
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run("Classification: ").bold = True
        meta.add_run(f"{entity.classification}\n")
        meta.add_run("Primary Key(s): ").bold = True
        meta.add_run(f"{', '.join(entity.primary_keys) if entity.primary_keys else 'None defined'}\n")
        meta.add_run("Attribute Count: ").bold = True
        meta.add_run(f"{entity.attribute_count}\n")
        
        # Source coverage
        sources = []
        if entity.source_coverage.get("guardrails"):
            sources.append("Guardrails")
        if entity.source_coverage.get("glue"):
            sources.append("Glue")
        if entity.source_coverage.get("ncpdp"):
            sources.append("NCPDP")
        if entity.source_coverage.get("fhir"):
            sources.append("FHIR")
        
        meta.add_run("Sources: ").bold = True
        meta.add_run(", ".join(sources) if sources else "None")
    
    doc.add_page_break()