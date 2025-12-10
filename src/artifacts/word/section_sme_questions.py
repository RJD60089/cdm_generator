# src/artifacts/word/section_sme_questions.py
"""SME Questions section for Word CDM document."""

from docx import Document
from typing import Optional
from src.artifacts.common.gap_extractor import GapExtractor


def add_sme_questions_section(doc: Document, gap_extractor: Optional[GapExtractor]):
    """Add SME Questions section."""
    
    doc.add_heading("7. SME Questions", level=1)
    
    doc.add_paragraph(
        "The following questions require subject matter expert input to finalize the CDM. "
        "Please provide responses in the 'Answer' column or discuss with the data governance team."
    )
    
    if not gap_extractor:
        doc.add_paragraph("No consolidation analysis available.")
        doc.add_page_break()
        return
    
    questions = gap_extractor.get_sme_questions()
    
    if not questions:
        doc.add_paragraph("No SME questions identified.")
        doc.add_page_break()
        return
    
    # Questions table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    headers = ["#", "Question", "Answer"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for idx, question in enumerate(questions, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = question.question_id or str(idx)
        row_cells[1].text = question.question_text or "-"
        row_cells[2].text = ""  # Empty for SME to fill in
    
    doc.add_paragraph()
    doc.add_paragraph(
        f"Total questions: {len(questions)}",
        style='Intense Quote'
    )
    
    doc.add_page_break()