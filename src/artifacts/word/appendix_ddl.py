# src/artifacts/word/appendix_ddl.py
"""DDL Appendix for Word CDM document."""

from docx import Document
from docx.shared import Pt
from pathlib import Path
from typing import Optional


def add_ddl_appendix(doc: Document, ddl_path: Optional[Path]):
    """Add DDL Script appendix."""
    
    doc.add_heading("Appendix A: DDL Script", level=1)
    
    if not ddl_path or not ddl_path.exists():
        doc.add_paragraph(
            "DDL script not available. Generate artifacts (Step 7) to create DDL."
        )
        doc.add_page_break()
        return
    
    doc.add_paragraph(f"Source file: {ddl_path.name}")
    doc.add_paragraph()
    
    # Read DDL content
    try:
        with open(ddl_path, 'r', encoding='utf-8') as f:
            ddl_content = f.read()
        
        # Add as code block (using monospace font)
        p = doc.add_paragraph()
        run = p.add_run(ddl_content)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        
    except Exception as e:
        doc.add_paragraph(f"Error reading DDL file: {e}")
    
    doc.add_page_break()
