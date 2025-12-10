# src/artifacts/word/section_executive_summary.py
"""Executive Summary section for Word CDM document."""

from docx import Document
from docx.shared import Pt, Inches
from typing import Optional

from src.artifacts.common.cdm_extractor import CDMExtractor
from src.artifacts.common.gap_extractor import GapExtractor


def add_executive_summary(
    doc: Document,
    extractor: CDMExtractor,
    gap_extractor: Optional[GapExtractor]
):
    """Add Executive Summary section."""
    
    doc.add_heading("1. Executive Summary", level=1)
    
    # Overview paragraph
    doc.add_paragraph(
        f"This document defines the Canonical Data Model (CDM) for the {extractor.domain} domain. "
        f"The CDM provides a standardized representation of {extractor.domain.lower()} data elements "
        f"to support data integration, analytics, and interoperability across systems."
    )
    
    # Key metrics table
    doc.add_heading("Key Metrics", level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    metrics = [
        ("Total Entities", str(extractor.entity_count)),
        ("Total Attributes", str(extractor.attribute_count)),
        ("Total Relationships", str(len(extractor.get_relationships()))),
    ]
    
    # Add gap metrics if available
    if gap_extractor:
        metrics.extend([
            ("Fields Requiring Review", str(len(gap_extractor.get_requires_review_fields()))),
            ("Unmapped Source Fields", str(len(gap_extractor.get_unmapped_fields()))),
            ("SME Questions", str(len(gap_extractor.get_sme_questions()))),
        ])
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    doc.add_paragraph()  # Spacing
    
    # Source coverage
    doc.add_heading("Source Coverage", level=2)
    
    coverage = extractor.get_source_coverage_summary()
    
    coverage_table = doc.add_table(rows=1, cols=2)
    coverage_table.style = 'Table Grid'
    
    hdr = coverage_table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Attributes Mapped"
    
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Order: Guardrails, Glue, NCPDP, FHIR
    for source in ["guardrails", "glue", "ncpdp", "fhir"]:
        count = coverage.get(source, 0)
        if count > 0:
            row = coverage_table.add_row().cells
            row[0].text = source.upper() if source in ["ncpdp", "fhir"] else source.title()
            row[1].text = str(count)
    
    doc.add_page_break()
