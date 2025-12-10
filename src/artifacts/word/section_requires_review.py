# src/artifacts/word/section_requires_review.py
"""Requires Review section for Word CDM document."""

from docx import Document
from typing import Optional
from src.artifacts.common.cdm_extractor import CDMExtractor
from src.artifacts.common.gap_extractor import GapExtractor


def add_requires_review_section(
    doc: Document,
    extractor: CDMExtractor,
    gap_extractor: Optional[GapExtractor]
):
    """Add Requires Review section."""
    
    doc.add_heading("6. Requires Review", level=1)
    
    doc.add_paragraph(
        "The following items have been flagged for SME review. These represent areas where "
        "additional clarification or decision-making is needed."
    )
    
    if not gap_extractor:
        doc.add_paragraph("No gap analysis available. Run Step 6 with gap analysis enabled.")
        doc.add_page_break()
        return
    
    review_items = gap_extractor.get_requires_review_fields()
    
    if not review_items:
        doc.add_paragraph("No items requiring review.")
        doc.add_page_break()
        return
    
    # Group by entity
    by_entity = {}
    for item in review_items:
        entity = item.cdm_entity or "Unknown"
        if entity not in by_entity:
            by_entity[entity] = []
        by_entity[entity].append(item)
    
    # Create table for each entity
    for entity, items in sorted(by_entity.items()):
        doc.add_heading(entity, level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header
        headers = ["Attribute", "Issue", "Source"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.cdm_attribute or "-"
            row_cells[1].text = item.review_reason or "-"
            row_cells[2].text = item.source_type or "-"
        
        doc.add_paragraph()  # Spacing
    
    doc.add_paragraph(
        f"Total items requiring review: {len(review_items)}",
        style='Intense Quote'
    )
    
    doc.add_page_break()