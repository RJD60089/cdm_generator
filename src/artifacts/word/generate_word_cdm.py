# src/artifacts/word/generate_word_cdm.py
"""
DDL Script Word Document Generator

Generates a simple Word document containing only the DDL script.
Format: Portrait, narrow margins, Calibri 10pt.
"""

from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config.config_parser import AppConfig
from src.artifacts.common.cdm_extractor import CDMExtractor
from src.artifacts.word.generate_ddl import generate_ddl_file, generate_ddl
from src.artifacts.word.generate_lucidchart_csv import generate_lucidchart_files


def setup_document(doc: Document):
    """Configure document: portrait, narrow margins, Calibri 10pt."""
    
    # Set narrow margins for all sections
    for section in doc.sections:
        # Portrait orientation (default)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Normal style - Calibri 10pt
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(14)
    h1.font.bold = True


def generate_word_ddl(
    config: AppConfig,
    cdm_path: Path,
    outdir: Path,
    dialect: str = "sqlserver",
    schema: str = "dbo"
) -> Path:
    """
    Generate Word document containing DDL script.
    
    Args:
        config: App configuration
        cdm_path: Path to Full CDM JSON
        outdir: Output directory
        dialect: SQL dialect (sqlserver, postgresql, mysql)
        schema: Database schema name
    
    Returns:
        Path to generated Word document
    """
    
    print(f"   Generating DDL Word document...")
    
    # Load CDM
    extractor = CDMExtractor(cdm_path=cdm_path)
    
    # Generate DDL content
    print(f"      Generating DDL ({dialect})...")
    ddl_content = generate_ddl(extractor, dialect=dialect, schema=schema)
    
    # Create document
    doc = Document()
    setup_document(doc)
    
    # Title
    title = doc.add_heading(f"{config.cdm.domain} CDM DDL Script", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Dialect: {dialect.upper()}\n")
    meta.add_run(f"Schema: {schema}\n")
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run(f"Entities: {extractor.entity_count}\n")
    
    doc.add_paragraph()  # spacer
    
    # DDL content
    ddl_para = doc.add_paragraph()
    ddl_run = ddl_para.add_run(ddl_content)
    ddl_run.font.name = 'Consolas'
    ddl_run.font.size = Pt(9)
    
    # Save
    artifacts_dir = outdir / "artifacts"
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    
    domain_safe = config.cdm.domain.lower().replace(' ', '_')
    output_file = artifacts_dir / f"{domain_safe}_{dialect}_ddl_script.docx"
    
    doc.save(str(output_file))
    
    print(f"   ✓ DDL document saved: {output_file.name}")
    
    return output_file


def generate_ddl_and_csv(
    config: AppConfig,
    cdm_path: Path,
    outdir: Path,
    dialect: str = "sqlserver",
    schema: str = "dbo"
) -> dict:
    """
    Generate DDL SQL file and LucidChart CSV file.
    
    Args:
        config: App configuration
        cdm_path: Path to Full CDM JSON
        outdir: Output directory
        dialect: SQL dialect
        schema: Database schema name
    
    Returns:
        Dict with paths to generated files
    """
    
    extractor = CDMExtractor(cdm_path=cdm_path)
    
    outputs = {}
    
    # Generate DDL file
    print(f"   Generating DDL file ({dialect})...")
    ddl_path = generate_ddl_file(
        extractor=extractor,
        outdir=outdir,
        domain=config.cdm.domain,
        dialect=dialect,
        schema=schema
    )
    outputs["ddl"] = ddl_path
    print(f"   ✓ DDL: {ddl_path.name}")
    
    # Generate LucidChart CSV from DDL
    print(f"   Generating LucidChart CSV...")
    csv_files = generate_lucidchart_files(
        ddl_path=ddl_path,
        outdir=outdir,
        domain=config.cdm.domain,
        dialect=dialect,
        schema=schema
    )
    outputs.update(csv_files)
    for name, path in csv_files.items():
        print(f"   ✓ {name}: {path.name}")
    
    return outputs


# Alias for backward compatibility
generate_word_cdm = generate_word_ddl
