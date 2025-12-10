# src/artifacts/word/appendix_lucidchart.py
"""LucidChart ERD Appendix for Word CDM document."""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Optional


def add_lucidchart_appendix(doc: Document, csv_path: Optional[Path] = None):
    """Add LucidChart ERD appendix placeholder."""
    
    doc.add_heading("Appendix B: Entity Relationship Diagram", level=1)
    
    doc.add_paragraph(
        "This section is reserved for the Entity Relationship Diagram (ERD). "
        "The ERD should be manually added after generating the diagram in LucidChart "
        "or a similar tool."
    )
    
    doc.add_paragraph()
    
    # CSV file reference
    if csv_path and csv_path.exists():
        doc.add_heading("LucidChart Import File", level=2)
        
        p = doc.add_paragraph()
        p.add_run("CSV file for import: ").bold = True
        p.add_run(f"{csv_path.name}\n")
        p.add_run("Location: ").bold = True
        p.add_run(f"{csv_path.parent}\n")
        
        doc.add_paragraph(
            "This CSV file can be imported into LucidChart to auto-generate the ERD. "
            "In LucidChart: File → Import Data → Import from CSV."
        )
        
        doc.add_paragraph()
    
    # Instructions
    doc.add_heading("Instructions", level=2)
    
    instructions = [
        "Import the CSV file into LucidChart (or create ERD manually)",
        "Arrange entities and relationships as needed",
        "Export the ERD as PDF or PNG",
        "Open this document in Microsoft Word",
        "Position cursor below this section",
        "Insert → Pictures → select the exported ERD image",
        "Resize as needed to fit the page"
    ]
    
    for idx, instruction in enumerate(instructions, 1):
        doc.add_paragraph(f"{idx}. {instruction}")
    
    doc.add_paragraph()
    
    # Placeholder area
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.add_run("\n\n[Insert ERD Image Here]\n\n").italic = True
    
    doc.add_page_break()