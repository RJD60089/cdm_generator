# src/artifacts/word/section_unmapped.py
"""Unmapped Fields section for Word CDM document."""

from docx import Document
from typing import Optional
from src.artifacts.common.gap_extractor import GapExtractor


def add_unmapped_section(doc: Document, gap_extractor: Optional[GapExtractor]):
    """Add Unmapped Fields section."""
    
    doc.add_heading("8. Unmapped Fields", level=1)
    
    doc.add_paragraph(
        "The following source fields were not mapped to the CDM. SME review is needed to determine "
        "if these fields should be added to the CDM, mapped to existing attributes, or excluded."
    )
    
    if not gap_extractor:
        doc.add_paragraph("No gap analysis available.")
        doc.add_page_break()
        return
    
    unmapped = gap_extractor.get_unmapped_fields()
    
    if not unmapped:
        doc.add_paragraph("All source fields have been mapped. No unmapped fields.")
        doc.add_page_break()
        return
    
    # Group by source
    by_source = {}
    for item in unmapped:
        source = item.source_type or "Unknown"
        if source not in by_source:
            by_source[source] = []
        by_source[source].append(item)
    
    # Create table for each source
    for source, items in sorted(by_source.items()):
        doc.add_heading(f"Source: {source.upper()}", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header
        headers = ["Source Entity", "Source Field", "Disposition"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.source_entity or "-"
            row_cells[1].text = item.source_attribute or "-"
            row_cells[2].text = ""  # Empty for SME decision
        
        doc.add_paragraph()  # Spacing
    
    doc.add_paragraph(
        f"Total unmapped fields: {len(unmapped)}",
        style='Intense Quote'
    )
    
    doc.add_page_break()