# src/artifacts/word/section_cdes.py
"""Critical Data Elements section for Word CDM document."""

from docx import Document
from docx.shared import Pt
from src.artifacts.common.cdm_extractor import CDMExtractor
from src.artifacts.common.cde_identifier import CDEIdentifier


def add_cdes_section(
    doc: Document,
    extractor: CDMExtractor,
    cde_identifier: CDEIdentifier
):
    """Add Critical Data Elements section."""
    
    doc.add_heading("5. Critical Data Elements", level=1)
    
    doc.add_paragraph(
        "Critical Data Elements (CDEs) are the most important data elements in the CDM. "
        "These are candidate selections based on business criticality criteria and require "
        "SME validation to confirm."
    )
    
    doc.add_heading("Selection Criteria", level=2)
    
    criteria = [
        "Required for core business transactions (claims routing, adjudication)",
        "Subject to external reporting, audit, or regulatory obligation",
        "Would break downstream systems or integrations if incorrect",
        "Contractual requirement with clients or trading partners",
        "Primary business identifier used across systems"
    ]
    
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    # CDE table
    doc.add_heading("Candidate CDEs", level=2)
    
    cdes = cde_identifier.identify_cdes()
    
    if not cdes:
        doc.add_paragraph("No Critical Data Elements identified. SME input required.")
        doc.add_page_break()
        return
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header
    headers = ["Entity", "Attribute", "Justification", "PII/PHI"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for cde in cdes:
        row_cells = table.add_row().cells
        row_cells[0].text = cde.entity_name
        row_cells[1].text = cde.attribute_name
        row_cells[2].text = cde.business_justification[:100] if len(cde.business_justification) > 100 else cde.business_justification
        
        flags = []
        if cde.is_pii:
            flags.append("PII")
        if cde.is_phi:
            flags.append("PHI")
        row_cells[3].text = ", ".join(flags) if flags else "-"
    
    doc.add_paragraph()
    doc.add_paragraph(
        f"Total CDEs: {len(cdes)}",
        style='Intense Quote'
    )
    
    doc.add_page_break()
